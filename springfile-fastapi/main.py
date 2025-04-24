import io
import asyncio # Import asyncio for sleep
import os
import re # For parsing GCS URI
import tempfile # For handling PDF loading from bytes
from typing import List # For type hinting

from fastapi import FastAPI, File, UploadFile, HTTPException, status # Added status
from fastapi.responses import StreamingResponse, JSONResponse # Added JSONResponse
from google.cloud import storage # Import GCS client
from pydantic import BaseModel # For request body validation
from docx import Document
from docx.shared import Pt

# Langchain specific imports
# UnstructuredWordDocumentLoader might not be needed if loading directly from bytes
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
import chromadb # Import ChromaDB client

app = FastAPI()

# --- Configuration ---
ALLOWED_EXTENSIONS = {".docx", ".pdf"}
# UPLOADS_DIR removed - no longer reading from local FS
MODEL_NAME = "BAAI/bge-small-zh-v1.5"
MODEL_KWARGS = {'device': 'cpu'} # Use CPU. Change to 'cuda' for GPU if available and configured.
ENCODE_KWARGS = {'normalize_embeddings': True} # Or False, depending on use case
CHROMA_DB_PATH = "./chroma_db" # Directory to store ChromaDB data
COLLECTION_NAME = "document_embeddings"

# --- Initialize Embedding Model (globally) ---
# This might take time on first run as the model downloads
try:
    print(f"Initializing BGE embedding model: {MODEL_NAME}...")
    embedding_model = HuggingFaceBgeEmbeddings(
        model_name=MODEL_NAME,
        model_kwargs=MODEL_KWARGS,
        encode_kwargs=ENCODE_KWARGS
    )
    print("BGE embedding model initialized successfully.")
except Exception as e:
    print(f"FATAL: Failed to initialize embedding model: {e}")
    embedding_model = None # Indicate failure

# --- Initialize GCS Client ---
try:
    print("Initializing Google Cloud Storage client...")
    storage_client = storage.Client()
    print("Google Cloud Storage client initialized successfully.")
except Exception as e:
    print(f"FATAL: Failed to initialize Google Cloud Storage client: {e}")
    storage_client = None

# --- Initialize ChromaDB Client and Collection ---
try:
    print(f"Initializing ChromaDB client at path: {CHROMA_DB_PATH}...")
    chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
    # Get or create the collection
    collection = chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        # Optional: Specify embedding function if you want Chroma to handle it,
        # but we are providing embeddings manually here.
        # metadata={"hnsw:space": "cosine"} # Example metadata, adjust as needed
    )
    print(f"ChromaDB collection '{COLLECTION_NAME}' ready.")
except Exception as e:
    print(f"FATAL: Failed to initialize ChromaDB: {e}")
    chroma_client = None
    collection = None

# --- Initialize Text Splitter ---
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50,
)

# --- Pydantic Models for Input ---
# Changed model to accept GCS URI
class GcsUriInput(BaseModel):
    gcs_uri: str # e.g., "gs://your-bucket-name/your-object-name.docx"

class SearchQueryInput(BaseModel):
    query: str # The text query for similarity search
    n_results: int = 5 # Number of results to return, default 5

@app.post("/preprocess/docx/")
async def preprocess_docx(file: UploadFile = File(...)):
    """
    Accepts a DOCX file, adds "hello world" to the beginning,
    and returns the modified DOCX file.
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only .docx files are accepted.")

    try:
        # --- Add 3-second delay for testing ---
        print("Processing started, waiting 3 seconds...")
        await asyncio.sleep(3)
        print("Delay finished, continuing processing.")
        # --------------------------------------

        # Read the uploaded file content into memory
        content = await file.read()
        file_stream = io.BytesIO(content)

        # Load the document using python-docx
        document = Document(file_stream)

        # Create a new paragraph with "hello world"
        # Insert it at the very beginning of the document body
        paragraph = document.add_paragraph()
        run = paragraph.add_run("hello world")
        run.font.size = Pt(12) # Optional: set font size

        # Move the new paragraph to the beginning
        # The Document object's body element contains paragraphs and tables.
        # We access the underlying XML element (body) and insert the new paragraph's XML element (p) at the start.
        body = document.element.body
        body.insert(0, paragraph._p) # _p gives the underlying lxml element

        # Save the modified document to a BytesIO stream
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0) # Reset stream position to the beginning

        # Return the modified file as a streaming response
        return StreamingResponse(
            output_stream,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="processed_{file.filename}"'}
        )

    except Exception as e:
        # Log the error for debugging
        print(f"Error processing file {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Could not process file: {e}")

@app.post("/embedding/")
async def embedding(payload: GcsUriInput): # Changed input model
    """
    Accepts a GCS URI (e.g., gs://bucket/object.docx),
    downloads the corresponding DOCX or PDF file from GCS, splits text,
    generates embeddings using the configured BGE model,
    stores them in ChromaDB, overwriting existing entries for the same GCS URI,
    and returns a success message.
    """
    if embedding_model is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Embedding model is not available.")
    if collection is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Vector database (ChromaDB) is not available.")
    if storage_client is None:
         raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="GCS client is not available.")

    gcs_uri = payload.gcs_uri
    print(f"Received request to embed file from GCS URI: {gcs_uri}")

    # --- GCS URI Parsing and Validation ---
    match = re.match(r"gs://([^/]+)/(.+)", gcs_uri)
    if not match:
        print(f"Invalid GCS URI format: {gcs_uri}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid GCS URI format. Expected gs://bucket-name/object-name"
        )

    bucket_name = match.group(1)
    object_name = match.group(2)
    print(f"Parsed GCS URI: Bucket='{bucket_name}', Object='{object_name}'")

    # Determine file extension from object name
    _, file_extension = os.path.splitext(object_name)
    file_extension = file_extension.lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        print(f"Invalid file type for embedding based on GCS object name: {file_extension}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file type based on GCS object name. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    try:
        # --- Download File Content from GCS ---
        print(f"Downloading content from gs://{bucket_name}/{object_name}...")
        try:
            bucket = storage_client.get_bucket(bucket_name)
            blob = bucket.blob(object_name)
            if not blob.exists():
                 print(f"File not found in GCS: gs://{bucket_name}/{object_name}")
                 raise HTTPException(
                     status_code=status.HTTP_404_NOT_FOUND,
                     detail=f"File not found in GCS at specified URI: {gcs_uri}"
                 )
            file_content_bytes = blob.download_as_bytes()
            print(f"Successfully downloaded {len(file_content_bytes)} bytes from GCS.")
        except Exception as e:
            print(f"Error downloading file from GCS: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to download file from GCS: {e}"
            )

        # --- Document Loading & Text Extraction (from downloaded bytes) ---
        all_text = ""
        print(f"Processing file extension: {file_extension}")
        if file_extension == ".docx":
            print(f"Loading content using python-docx from downloaded bytes...")
            try:
                document = Document(io.BytesIO(file_content_bytes)) # Load from BytesIO
                all_text = "\n\n".join([para.text for para in document.paragraphs])
                print(f"Successfully extracted text from DOCX bytes (length: {len(all_text)}).")
            except Exception as e:
                print(f"Error processing DOCX file content from bytes: {e}")
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Error processing DOCX file content from bytes: {e}"
                )
        elif file_extension == ".pdf":
             # UnstructuredPDFLoader needs a file path. Save bytes to a temporary file.
            print(f"Loading content using UnstructuredPDFLoader via temporary file...")
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(file_content_bytes)
                    temp_pdf_path = temp_pdf.name
                    print(f"Saved PDF bytes to temporary file: {temp_pdf_path}")

                loader = UnstructuredPDFLoader(temp_pdf_path)
                documents = loader.load()
                all_text = "\n\n".join([doc.page_content for doc in documents])
                print(f"Successfully loaded {len(documents)} PDF document parts via temp file (total text length: {len(all_text)}).")

            except Exception as e:
                print(f"Error processing PDF file content via temporary file: {e}")
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Error processing PDF file content via temporary file: {e}"
                )
            finally:
                # Clean up the temporary file
                if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
                    try:
                        os.remove(temp_pdf_path)
                        print(f"Cleaned up temporary file: {temp_pdf_path}")
                    except Exception as cleanup_e:
                        print(f"Warning: Failed to clean up temporary file {temp_pdf_path}: {cleanup_e}")
        else:
             # Should not happen due to earlier check
             print(f"Internal error: File type processing failed for extension {file_extension}")
             raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal error: File type processing failed.")

        # --- Text Splitting ---
        if not all_text.strip():
             print("Warning: Loaded document content is empty or whitespace only.")
             # Use gcs_uri in response
             return JSONResponse(content={"gcs_uri": gcs_uri, "embeddings": []})

        print(f"Splitting text (total length: {len(all_text)})...")
        chunks = text_splitter.split_text(all_text)
        print(f"Split text into {len(chunks)} chunks.")

        if not chunks:
             print("Warning: Text splitting resulted in zero chunks.")
             # Use gcs_uri in response
             return JSONResponse(content={"gcs_uri": gcs_uri, "embeddings": []})

        # --- Embedding Generation ---
        print(f"Generating embeddings for {len(chunks)} chunks using {MODEL_NAME}...")
        try:
            embeddings: List[List[float]] = embedding_model.embed_documents(chunks)
            print(f"Successfully generated {len(embeddings)} embeddings (dimension: {len(embeddings[0]) if embeddings else 'N/A'}).")
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate embeddings: {e}"
            )

        # --- Store Embeddings in ChromaDB ---
        if embeddings:
            # Create unique IDs for each chunk based on GCS URI and index
            ids = [f"{gcs_uri}_{i}" for i in range(len(chunks))]
            # Create metadata (optional, but useful) - store GCS URI as source
            metadatas = [{"source": gcs_uri, "chunk_index": i} for i in range(len(chunks))]

            try:
                # 1. Delete existing embeddings for this GCS URI
                print(f"Deleting existing embeddings for GCS URI: {gcs_uri}...")
                # Query for existing IDs associated with this GCS URI
                existing_docs = collection.get(where={"source": gcs_uri}, include=[]) # Only need IDs
                if existing_docs and existing_docs['ids']:
                    print(f"Found {len(existing_docs['ids'])} existing embeddings to delete.")
                    collection.delete(ids=existing_docs['ids'])
                    print("Existing embeddings deleted.")
                else:
                    print("No existing embeddings found for this GCS URI.")

                # 2. Add new embeddings
                print(f"Adding {len(chunks)} new embeddings to ChromaDB collection '{COLLECTION_NAME}'...")
                collection.add(
                    embeddings=embeddings,
                    documents=chunks, # Store the text chunk itself
                    metadatas=metadatas,
                    ids=ids
                )
                print("New embeddings added successfully.")

            except Exception as e:
                print(f"Error interacting with ChromaDB: {e}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to store embeddings in vector database: {e}"
                )
        else:
            print("No embeddings generated, skipping database storage.")


        # --- Response ---
        # Return success message using GCS URI
        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={
                "message": f"Successfully processed and stored embeddings for GCS URI: {gcs_uri}",
                "gcs_uri": gcs_uri,
                "chunks_processed": len(chunks),
                "embeddings_stored": len(embeddings) if embeddings else 0
            }
        )

    except HTTPException as e:
        # Re-raise HTTPExceptions
        print(f"HTTPException occurred: {e.status_code} - {e.detail}")
        raise e
    except Exception as e:
        # Catch other unexpected errors
        print(f"An unexpected error occurred during embedding: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An unexpected error occurred: {e}"
        )

@app.post("/search/")
async def search_embeddings(payload: SearchQueryInput):
    """
    Accepts a text query, generates its embedding, and searches the ChromaDB
    collection for the most similar document chunks.
    Returns the top N results (default 5).
    """
    if embedding_model is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Embedding model is not available.")
    if collection is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Vector database (ChromaDB) is not available.")

    query_text = payload.query
    n_results = payload.n_results
    print(f"Received search query: '{query_text}', requesting {n_results} results.")

    if not query_text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Query text cannot be empty.")
    if not 1 <= n_results <= 50: # Add reasonable limit
         raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Number of results must be between 1 and 50.")

    try:
        # --- Generate Query Embedding ---
        print(f"Generating embedding for query: '{query_text}'...")
        query_embedding = embedding_model.embed_query(query_text)
        print("Query embedding generated successfully.")

        # --- Query ChromaDB ---
        print(f"Querying ChromaDB collection '{COLLECTION_NAME}' for {n_results} similar documents...")
        results = collection.query(
            query_embeddings=[query_embedding], # Note: query expects a list of embeddings
            n_results=n_results,
            include=['documents', 'metadatas', 'distances'] # Include relevant info
        )
        print(f"ChromaDB query completed. Found results: {results}")

        # --- Format and Return Results ---
        # Results structure example:
        # {'ids': [['id1', 'id2']], 'distances': [[0.1, 0.2]], 'metadatas': [[{'source': 'file1'}, {'source': 'file2'}]], 'embeddings': None, 'documents': [['chunk1 text', 'chunk2 text']]}
        # We need to extract the relevant parts for our response.
        # Since we query with one embedding, we access the first element of each list.
        formatted_results = []
        if results and results.get('ids') and results['ids'][0]:
            ids = results['ids'][0]
            distances = results['distances'][0]
            metadatas = results['metadatas'][0]
            documents = results['documents'][0]

            for i in range(len(ids)):
                formatted_results.append({
                    "id": ids[i],
                    "distance": distances[i],
                    "metadata": metadatas[i],
                    "document": documents[i]
                })

        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={
                "query": query_text,
                "results": formatted_results
            }
        )

    except HTTPException as e:
        # Re-raise HTTPExceptions
        print(f"HTTPException occurred during search: {e.status_code} - {e.detail}")
        raise e
    except Exception as e:
        # Catch other unexpected errors
        print(f"An unexpected error occurred during search: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An unexpected error occurred during search: {e}"
        )


if __name__ == "__main__":
    import uvicorn
    import os
    # Read port from environment variable PORT, default to 8001 for local dev
    port = int(os.environ.get("PORT", 8001))
    # Run with: uvicorn main:app --reload --port $PORT (or default 8001)
    # Includes endpoints: /preprocess/docx/, /embedding/, /search/
    print(f"Starting Uvicorn server on port {port}...")
    # Use reload=False for production/container environment
    uvicorn.run(app, host="0.0.0.0", port=port, reload=False)
