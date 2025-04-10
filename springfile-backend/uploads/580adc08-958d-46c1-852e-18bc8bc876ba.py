from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn

def is_page_number(text):
    stripped = text.strip(' -')
    return stripped.isdigit()

def remove_unwanted_elements(input_path, output_path):
    doc = Document(input_path)
    
    # Identify section breaks (paragraphs containing sectPr)
    section_breaks = []
    for idx, para in enumerate(doc.paragraphs):
        p_pr = para._element.pPr
        if p_pr is not None and p_pr.find(qn('w:sectPr')) is not None:
            section_breaks.append(idx)
    
    section_break_indices = set(section_breaks)
    
    # Split document into sections
    sections = []
    prev = 0
    for br in section_breaks:
        sections.append((prev, br))
        prev = br + 1
    sections.append((prev, len(doc.paragraphs) - 1))
    
    deletion_indices = set()
    
    for start, end in sections:
        # Remove right-aligned text in first two paragraphs
        for i in range(start, min(start + 2, end + 1)):
            if i >= len(doc.paragraphs) or i in section_break_indices:
                continue
            para = doc.paragraphs[i]
            if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                deletion_indices.add(i)
        
        # Remove page numbers in last two paragraphs
        for i in range(max(end - 1, start), end + 1):
            if i >= len(doc.paragraphs) or i in section_break_indices:
                continue
            para = doc.paragraphs[i]
            if is_page_number(para.text):
                deletion_indices.add(i)
                # Check for trailing empty paragraph
                if i + 1 < len(doc.paragraphs) and (i + 1) not in section_break_indices:
                    next_para = doc.paragraphs[i + 1]
                    if next_para.text.strip() == '':
                        deletion_indices.add(i + 1)
    
    # Delete collected paragraphs in reverse order
    for index in sorted(deletion_indices, reverse=True):
        if index < len(doc.paragraphs):
            para = doc.paragraphs[index]
            para_element = para._element
            para_element.getparent().remove(para_element)
    
    doc.save(output_path)

# Usage
remove_unwanted_elements("/home/disky/lab/股权激励计划/3.docx", "/home/disky/lab/股权激励计划/3_1.docx")